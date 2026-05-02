"""
Pure draft export helpers: JSON, CSV, Markdown, Markdown+ZIP, Word (.docx) — no FastAPI imports.
Plain Markdown omits inline base64 images (readable .md); JSON/CSV retain full data URIs;
`mdzip` bundles questions.md + images/; DOCX embeds PNG/JPEG.
"""

from __future__ import annotations

import base64
import binascii
import csv
import io
import json
import re
import zipfile
from pathlib import Path
from typing import Any

SUPPORTED_FORMATS = frozenset({"json", "csv", "md", "markdown", "docx", "mdzip"})

MEDIA_TYPES = {
    "json": "application/json",
    "csv": "text/csv",
    "md": "text/markdown",
    "markdown": "text/markdown",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "mdzip": "application/zip",
}

_CSV_HEADER = [
    "question_number",
    "text",
    "option_a",
    "option_b",
    "option_c",
    "option_d",
    "correct_answer",
    "core_concept",
    "detailed_explanation",
    "option_analysis_a",
    "option_analysis_b",
    "option_analysis_c",
    "option_analysis_d",
    "flashcards_count",
    "diagram_data_uri",
]


def _canonical_fmt(fmt: str) -> str:
    f = fmt.lower()
    if f == "markdown":
        return "md"
    return f


def safe_filename(meta: dict[str, Any], fmt: str) -> str:
    """Build a filesystem-safe attachment name from draft metadata."""
    f = fmt.lower()
    canonical = "md" if f == "markdown" else f
    if canonical not in ("json", "csv", "md", "docx", "mdzip"):
        canonical = "json"

    paper = str(meta.get("paper_code") or "").strip()
    fname = str(meta.get("filename") or "").strip()
    draft_id = str(meta.get("id") or "draft").strip()
    stem = paper or (Path(fname).stem if fname else "") or draft_id
    stem = re.sub(r"[^\w\-.]+", "_", stem, flags=re.ASCII)
    stem = stem.strip("._") or "export"
    if len(stem) > 100:
        stem = stem[:100]

    if canonical == "mdzip":
        ext = "zip"
    elif canonical == "md":
        ext = "md"
    else:
        ext = canonical
    return f"{stem}.{ext}"


def to_json_bytes(meta: dict[str, Any]) -> bytes:
    return json.dumps(meta, indent=2, ensure_ascii=False).encode("utf-8")


def to_csv_bytes(meta: dict[str, Any]) -> bytes:
    buf = io.StringIO(newline="")
    writer = csv.writer(buf, quoting=csv.QUOTE_ALL, lineterminator="\n")
    writer.writerow(_CSV_HEADER)
    questions = meta.get("questions") or []
    for q in questions:
        opts = q.get("options") or {}
        if not isinstance(opts, dict):
            opts = {}
        gen = q.get("generated_content") or {}
        if not isinstance(gen, dict):
            gen = {}
        oa = gen.get("option_analysis") or {}
        if not isinstance(oa, dict):
            oa = {}
        flashcards = gen.get("flashcards") or []
        fc_count = len(flashcards) if isinstance(flashcards, list) else 0
        diagram = q.get("diagram_base64")
        if diagram is None:
            diagram = ""
        else:
            diagram = str(diagram)

        writer.writerow(
            [
                q.get("question_number", ""),
                q.get("text") or "",
                opts.get("A", ""),
                opts.get("B", ""),
                opts.get("C", ""),
                opts.get("D", ""),
                q.get("correct_answer") or "",
                gen.get("core_concept") or "",
                gen.get("detailed_explanation") or "",
                oa.get("A", ""),
                oa.get("B", ""),
                oa.get("C", ""),
                oa.get("D", ""),
                fc_count,
                diagram,
            ]
        )
    return buf.getvalue().encode("utf-8")


def _md_escape_pipe(s: str) -> str:
    return str(s).replace("|", "\\|")


def to_markdown_bytes(meta: dict[str, Any], *, image_paths: dict[str, str] | None = None) -> bytes:
    """
    Build Markdown. Plain export (no image_paths): diagrams are not inlined as base64;
    a short blockquote points users to mdzip/JSON/DOCX. With image_paths (from mdzip):
    emit ![...](images/Qn.png) relative links.
    """
    lines: list[str] = []
    lines.append("# Edmate draft export")
    lines.append("")
    lines.append(f"- **Subject:** {meta.get('subject', '')}")
    lines.append(f"- **Paper code:** {meta.get('paper_code', '')}")
    lines.append(f"- **Curriculum:** {meta.get('curriculum', '')}")
    lines.append(f"- **Draft id:** {meta.get('id', '')}")
    lines.append(f"- **Source file:** {meta.get('filename', '')}")
    ts = meta.get("timestamp") or meta.get("last_updated_at") or ""
    if ts:
        lines.append(f"- **Timestamp:** {ts}")
    lines.append("")
    lines.append("---")
    lines.append("")

    questions = meta.get("questions") or []
    for q in questions:
        qn = q.get("question_number", "?")
        qn_key = str(qn)
        lines.append(f"## Question {qn}")
        lines.append("")
        text = (q.get("text") or "").strip()
        if text:
            lines.append(text)
            lines.append("")

        diagram = q.get("diagram_base64")
        if diagram:
            dstr = str(diagram).strip()
            if image_paths and qn_key in image_paths:
                rel = image_paths[qn_key]
                lines.append(f"![Diagram for Q{qn}]({rel})")
                lines.append("")
            elif dstr.startswith("data:image/"):
                lines.append(
                    f"> Diagram for Q{qn} omitted in plain Markdown — use **Markdown+ZIP** "
                    f"(`format=mdzip`), **JSON**, or **Word** (`.docx`) export to retain the image."
                )
                lines.append("")
            else:
                lines.append(
                    f"> Diagram for Q{qn} omitted in plain Markdown — use Markdown+ZIP, JSON, or Word export."
                )
                lines.append("")

        opts = q.get("options") or {}
        if not isinstance(opts, dict):
            opts = {}
        correct = str(q.get("correct_answer") or "").strip().upper()

        lines.append("### Options")
        lines.append("")
        for letter in ("A", "B", "C", "D"):
            raw = opts.get(letter, "")
            label = f"**{letter}.**" if letter == correct else f"{letter}."
            lines.append(f"- {label} {_md_escape_pipe(str(raw))}")
        lines.append("")

        gen = q.get("generated_content") or {}
        if not isinstance(gen, dict):
            gen = {}

        cc = (gen.get("core_concept") or "").strip()
        if cc:
            lines.append("### Core concept")
            lines.append("")
            lines.append(cc)
            lines.append("")

        de = (gen.get("detailed_explanation") or "").strip()
        if de:
            lines.append("### Explanation")
            lines.append("")
            lines.append(de)
            lines.append("")

        oa = gen.get("option_analysis") or {}
        if isinstance(oa, dict) and any(str(oa.get(k, "")).strip() for k in ("A", "B", "C", "D")):
            lines.append("### Per-option analysis")
            lines.append("")
            for letter in ("A", "B", "C", "D"):
                chunk = str(oa.get(letter, "")).strip()
                if chunk:
                    lines.append(f"- **{letter}:** {chunk}")
            lines.append("")

        flashcards = gen.get("flashcards") or []
        if isinstance(flashcards, list) and flashcards:
            lines.append("### Flashcards")
            lines.append("")
            for i, fc in enumerate(flashcards, 1):
                if not isinstance(fc, dict):
                    continue
                fq = str(fc.get("question", "")).strip()
                fa = str(fc.get("answer", "")).strip()
                if fq or fa:
                    lines.append(f"{i}. **Q:** {fq}")
                    lines.append(f"   **A:** {fa}")
            lines.append("")

        lines.append("---")
        lines.append("")

    return "\n".join(lines).encode("utf-8")


def _diagram_data_uri_to_stream(uri: str) -> io.BytesIO | None:
    """Decode PNG/JPEG data URI to BytesIO for python-docx, or None if unsupported."""
    if not uri or not isinstance(uri, str):
        return None
    uri = uri.strip()
    m = re.match(r"data:image/(png|jpeg|jpg);base64,(.+)", uri, re.IGNORECASE | re.DOTALL)
    if not m:
        return None
    try:
        raw = base64.b64decode(m.group(2).strip(), validate=False)
    except (ValueError, binascii.Error):
        return None
    if not raw:
        return None
    bio = io.BytesIO(raw)
    bio.seek(0)
    return bio


def _data_uri_image_bytes_and_ext(uri: str) -> tuple[bytes, str] | None:
    """Return (raw_image_bytes, 'png'|'jpg') for PNG/JPEG data URIs, or None."""
    stream = _diagram_data_uri_to_stream(uri)
    if not stream:
        return None
    raw = stream.getvalue()
    m = re.match(r"data:image/(png|jpeg|jpg);base64,", uri.strip(), re.IGNORECASE)
    if not m:
        return raw, "png"
    g = m.group(1).lower()
    ext = "jpg" if g in ("jpeg", "jpg") else "png"
    return raw, ext


_ZIP_README = """Edmate Markdown bundle

- questions.md — question text, options, and generated explanations
- images/ — PNG or JPEG diagrams (when the draft had embeddable images)

Unpack the entire folder and open questions.md in any Markdown viewer; image links are relative.
"""


def to_markdown_zip_bytes(meta: dict[str, Any]) -> bytes:
    """ZIP containing questions.md, README.txt, and images/Q{n}.png|jpg for each diagram."""
    image_paths: dict[str, str] = {}
    image_payloads: list[tuple[str, bytes]] = []

    for q in meta.get("questions") or []:
        qn = q.get("question_number", "?")
        qn_key = str(qn)
        diag = q.get("diagram_base64")
        if not diag:
            continue
        dec = _data_uri_image_bytes_and_ext(str(diag))
        if not dec:
            continue
        raw, ext = dec
        arc = f"images/Q{qn}.{ext}"
        image_paths[qn_key] = arc
        image_payloads.append((arc, raw))

    md_body = to_markdown_bytes(meta, image_paths=image_paths if image_paths else None)
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("questions.md", md_body)
        zf.writestr("README.txt", _ZIP_README.encode("utf-8"))
        for arc, raw in image_payloads:
            zf.writestr(arc, raw)
    return out.getvalue()


def _xml_safe_text(value: Any) -> str:
    """Strip NUL and XML-illegal control chars so python-docx/lxml does not raise."""
    if value is None:
        return ""
    parts: list[str] = []
    for ch in str(value):
        o = ord(ch)
        if o == 0:
            continue
        if o < 32:
            if ch in "\t\n\r":
                parts.append(ch)
            continue
        if o in (0xFFFE, 0xFFFF):
            continue
        parts.append(ch)
    return "".join(parts)


def to_docx_bytes(meta: dict[str, Any]) -> bytes:
    """Build a Word document (.docx) from draft metadata."""
    from docx import Document  # type: ignore[import-untyped]
    from docx.shared import Inches  # type: ignore[import-untyped]

    doc = Document()
    doc.add_heading("Edmate draft export", 0)
    doc.add_paragraph(f"Subject: {_xml_safe_text(meta.get('subject', ''))}")
    doc.add_paragraph(f"Paper code: {_xml_safe_text(meta.get('paper_code', ''))}")
    doc.add_paragraph(f"Curriculum: {_xml_safe_text(meta.get('curriculum', ''))}")
    doc.add_paragraph(f"Draft id: {_xml_safe_text(meta.get('id', ''))}")
    doc.add_paragraph(f"Source file: {_xml_safe_text(meta.get('filename', ''))}")
    ts = meta.get("timestamp") or meta.get("last_updated_at") or ""
    if ts:
        doc.add_paragraph(f"Timestamp: {_xml_safe_text(ts)}")

    questions = meta.get("questions") or []
    for q in questions:
        qn = q.get("question_number", "?")
        doc.add_heading(f"Question {qn}", level=1)
        text = _xml_safe_text((q.get("text") or "").strip())
        if text:
            for line in text.split("\n"):
                doc.add_paragraph(_xml_safe_text(line))

        diagram = q.get("diagram_base64")
        if diagram:
            stream = _diagram_data_uri_to_stream(str(diagram))
            if stream is not None:
                try:
                    doc.add_picture(stream, width=Inches(5.5))
                except Exception:
                    doc.add_paragraph(
                        "(Diagram could not be embedded in Word — use JSON export for the full data URI.)"
                    )
            else:
                doc.add_paragraph(
                    "(Diagram is not PNG/JPEG — embed skipped; full URI is in JSON/CSV export.)"
                )

        opts = q.get("options") or {}
        if not isinstance(opts, dict):
            opts = {}
        correct = str(q.get("correct_answer") or "").strip().upper()
        if any(str(opts.get(L, "")).strip() for L in ("A", "B", "C", "D")):
            doc.add_heading("Options", level=2)
            for letter in ("A", "B", "C", "D"):
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(f"{letter}. {_xml_safe_text(opts.get(letter, ''))}")
                if letter == correct:
                    run.bold = True

        gen = q.get("generated_content") or {}
        if not isinstance(gen, dict):
            gen = {}

        cc = _xml_safe_text((gen.get("core_concept") or "").strip())
        if cc:
            doc.add_heading("Core concept", level=2)
            for line in cc.split("\n"):
                doc.add_paragraph(_xml_safe_text(line))

        de = _xml_safe_text((gen.get("detailed_explanation") or "").strip())
        if de:
            doc.add_heading("Explanation", level=2)
            for line in de.split("\n"):
                doc.add_paragraph(_xml_safe_text(line))

        oa = gen.get("option_analysis") or {}
        if isinstance(oa, dict) and any(str(oa.get(k, "")).strip() for k in ("A", "B", "C", "D")):
            doc.add_heading("Per-option analysis", level=2)
            for letter in ("A", "B", "C", "D"):
                chunk = _xml_safe_text(str(oa.get(letter, "")).strip())
                if chunk:
                    p = doc.add_paragraph()
                    p.add_run(f"{letter}. ").bold = True
                    p.add_run(chunk)

        flashcards = gen.get("flashcards") or []
        if isinstance(flashcards, list) and flashcards:
            doc.add_heading("Flashcards", level=2)
            for i, fc in enumerate(flashcards, 1):
                if not isinstance(fc, dict):
                    continue
                fq = _xml_safe_text(str(fc.get("question", "")).strip())
                fa = _xml_safe_text(str(fc.get("answer", "")).strip())
                if fq or fa:
                    doc.add_paragraph(f"{i}. Q: {fq}", style="List Bullet")
                    doc.add_paragraph(f"   A: {fa}", style="List Bullet")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def render(meta: dict[str, Any], fmt: str) -> bytes:
    """Serialize metadata to bytes for the given export format."""
    f = fmt.lower()
    if f not in SUPPORTED_FORMATS:
        raise ValueError(f"Unsupported format: {fmt}")
    canon = _canonical_fmt(f)
    if canon == "json":
        return to_json_bytes(meta)
    if canon == "csv":
        return to_csv_bytes(meta)
    if canon == "docx":
        return to_docx_bytes(meta)
    if canon == "mdzip":
        return to_markdown_zip_bytes(meta)
    return to_markdown_bytes(meta)
